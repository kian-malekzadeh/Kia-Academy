#!/usr/bin/env python3
"""Generate Kia Academy / Kia Academy comprehensive developer guidebook (.docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path("/workspace/docs/Kia_Academy_Project_Guidebook.docx")


def set_run_font(run, *, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x2F, 0x44)
    return h


def add_p(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, size=size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items, *, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * level)
        run = p.add_run(item)
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9, color=RGBColor(0x0A, 0x2F, 0x44))
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            run = cells[c_i].paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def toc_entry(doc, title, page_hint=""):
    p = doc.add_paragraph()
    run = p.add_run(title)
    set_run_font(run, size=11)
    if page_hint:
        run2 = p.add_run(f"  —  {page_hint}")
        set_run_font(run2, size=10, italic=True, color=RGBColor(0x5D, 0x7F, 0x95))


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ── Cover ──────────────────────────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("کیا آکادمی (Kia Academy)")
    set_run_font(r, bold=True, size=28, color=RGBColor(0x0A, 0x2F, 0x44))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Comprehensive Project Guidebook")
    set_run_font(r, bold=True, size=18, color=RGBColor(0x1F, 0x6E, 0x8C))

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("for New Developers")
    set_run_font(r, size=14, color=RGBColor(0x38, 0x5E, 0x77))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Monorepo: Kia Academy (Persian-first rebuild)\n"
        "Stack: Next.js 15 · NestJS 11 · Prisma · PostgreSQL 16\n"
        "Audience: engineers onboarding to apps/web, apps/api, packages/shared"
    )
    set_run_font(r, size=10, italic=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "This document inventories purpose, architecture, every major feature "
        "(including UI controls), roles, API, security, colors, fonts, algorithms, "
        "and development practices as implemented in the repository."
    )
    set_run_font(r, size=10)

    doc.add_page_break()

    # ── TOC ────────────────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", 1)
    sections_toc = [
        ("1. Project Overview & Purpose", "What the product is and why it exists"),
        ("2. Detailed Project Presentation", "Deep walkthrough of product doors and flows"),
        ("3. Architecture & Repository Layout", "Monorepo packages and request path"),
        ("4. Languages & Core Technologies", "Runtime languages and platform tech"),
        ("5. Tools, Libraries, Frameworks & Tooling", "Every dependency and why"),
        ("6. Client Roles & Access Control", "Guest, LEARNER, ADMIN, SUPER_ADMIN"),
        ("7. Website Sections & Routes", "Every page/route the app exposes"),
        ("8. Complete Feature Catalog", "Exhaustive UI/feature inventory"),
        ("9. Color Palette", "Requested hexes + every detected palette"),
        ("10. Typography & Fonts", "Site fonts and Material Studio catalog"),
        ("11. Algorithms & Data Structures", "Scoring, OTP, exam, roadmap logic"),
        ("12. API Surface", "Endpoints, auth flow, payments"),
        ("13. Security Measures", "Auth, guards, validation, headers"),
        ("14. Database Models (Prisma)", "Enums and entities"),
        ("15. Payments, Currency & Email", "IRR/Toman, gateways, EmailLog"),
        ("16. Internationalization & Theming", "fa RTL default, locales, light/dark"),
        ("17. Development Setup & Commands", "Local env, Postgres, scripts"),
        ("18. Testing & CI/CD", "Vitest, Jest, Playwright, GitHub Actions"),
        ("19. Seed Data, Demo Mode & Ops Notes", "Accounts and deployment caveats"),
        ("20. Glossary & Quick Reference", "Terms and file-path index"),
    ]
    for title_s, hint in sections_toc:
        toc_entry(doc, title_s, hint)

    doc.add_page_break()

    # ── 1. Overview ────────────────────────────────────────────────────────
    add_heading(doc, "1. Project Overview & Purpose", 1)
    add_p(
        doc,
        "کیا آکادمی (Kia Academy) is a Persian-first adaptive learning platform aimed at learners in Iran. "
        "The repository is organized around this product shape "
        "(see docs/REBUILD_ARCHITECTURE.md).",
    )
    add_heading(doc, "1.1 Main purpose", 2)
    add_bullets(
        doc,
        [
            "Give guests a minimal landing with two clear doors: free Material Studio (design tools) and Education (learning path).",
            "Onboard learners via Iranian mobile phone OTP (not email-first), then collect a profile.",
            "Run a free goal assessment (wizard) that personalizes a learning roadmap.",
            "Immediately offer a free readiness / preparations exam; show a scorecard only after submit.",
            "Sell roadmap bundles or individual courses (amounts stored in IRR; UI shows تومان).",
            "Deliver lessons in a Kia Learn player (sidebar, video, notes, HTML/CSS/JS playground).",
            "Give staff (ADMIN / SUPER_ADMIN) a permissioned admin panel for catalog, users, payments, and settings.",
        ],
    )
    add_heading(doc, "1.2 Business model (brief)", 2)
    add_p(
        doc,
        "Acquisition is freemium: Material Studio and the assessment + readiness exam are free. "
        "Monetization is course/roadmap bundle checkout through configurable gateways "
        "(dev sandbox, Zarinpal, IDPay, or Stripe). Default currency is IRR.",
    )
    add_heading(doc, "1.3 Live / demo surfaces", 2)
    add_bullets(
        doc,
        [
            "Full local stack: Next.js (3000) proxies /api → NestJS (3001) + PostgreSQL 16.",
            "GitHub Pages static export can run in NEXT_PUBLIC_DEMO_MODE with in-browser mocks (no hosted API).",
            "Homepage (Pages): https://kian-malekzadeh.github.io/Kia-Academy/ (rename GitHub repo to kia-academy for /kia-academy/)",
        ],
    )

    # ── 2. Detailed presentation ───────────────────────────────────────────
    add_heading(doc, "2. Detailed Project Presentation", 1)
    add_p(
        doc,
        "This section explains the product at a level useful for someone who has never opened the repo.",
    )
    add_heading(doc, "2.1 Two product doors", 2)
    add_p(
        doc,
        "Door A — Material Studio (/material): a client-only design workshop (color palettes, icons, animations, "
        "font/style tools). No NestJS dependency. State can persist in localStorage. Marketed as free.",
    )
    add_p(
        doc,
        "Door B — Education (/education): phone OTP → profile → assessment → readiness exam → roadmap → "
        "checkout → learn. Site TopBar/Footer appear only after profileComplete; guests rely on landing CTAs.",
    )
    add_heading(doc, "2.2 Canonical learner journey", 2)
    add_numbered(
        doc,
        [
            "Land on / (Persian hero: brand کیا آکادمی + one headline + Material / Education CTAs).",
            "Enter /education → enter Iranian phone → receive OTP (devCode in non-prod) → verify.",
            "Complete profile: firstName, lastName, city, email (phone read-only).",
            "Start assessment wizard (6 stages: goal, skills, personality, interests, learning style, hours/week).",
            "System saves Assessment + generates Roadmap; starts free timed readiness exam (ExamPlayer).",
            "During the exam: no green/red correctness feedback; autosave answers; 30-minute timer.",
            "Results page shows domain scores, pass/fail, roadmap mutations (unlocks / refreshers).",
            "Continue to /roadmap?roadmapId=… (missing id causes a misleading “complete assessment” error).",
            "Purchase ROADMAP_BUNDLE or COURSE(s) via /checkout → optional gateway UX → entitlements.",
            "Open /courses → /learn/[courseSlug]/[lessonSlug] for video, markdown, notes, playground.",
            "Optional: /bootcamp challenge (e.g. FizzBuzz) → /rewards unlocks.",
        ],
    )
    add_heading(doc, "2.3 Staff journey", 2)
    add_p(
        doc,
        "Staff use /login with email+password (seeded accounts). SUPER_ADMIN reaches full /admin. "
        "ADMIN users see only sidebar items allowed by their per-user adminPanelAccess matrix "
        "(defaults come from site settings adminAccess template when promoted).",
    )
    add_heading(doc, "2.4 What “adaptive” means here", 2)
    add_p(
        doc,
        "Adaptation is rule-based, not ML: skill answers map to a level band; primary interest selects a track; "
        "exam outcomes unlock the first module, may promote absolute beginners who score high, or prepend "
        "refresher modules for weak domains. Pricing sums MODULE_PRICES with a configurable bundle discount.",
    )
    add_heading(doc, "2.5 Important product caveats", 2)
    add_bullets(
        doc,
        [
            "Education always starts phone-first; do not auto-skip OTP for incomplete registrations.",
            "Live readiness UI must not reveal correctness until /readiness/results.",
            "Legacy readiness components (file-drop, flowchart, etc.) still exist under components/readiness/ "
            "but the live product path is the server-graded ExamPlayer.",
            "Catalog prices are IRR integers; Persian UI divides by 10 to show تومان.",
        ],
    )

    # ── 3. Architecture ────────────────────────────────────────────────────
    add_heading(doc, "3. Architecture & Repository Layout", 1)
    add_heading(doc, "3.1 High-level diagram", 2)
    add_p(
        doc,
        "Browser → apps/web (Next.js 15 App Router, React 19, CSS variables, next/font)\n"
        "         → rewrite /api/* → apps/api (NestJS 11, Passport JWT, Throttler, Helmet)\n"
        "         → Prisma 6 → PostgreSQL 16\n"
        "packages/shared ← types, tracks, exam bank, grading, entitlements (built to dist/ first)",
        italic=True,
    )
    add_heading(doc, "3.2 Monorepo packages", 2)
    add_table(
        doc,
        ["Package", "Path", "Role"],
        [
            ["kia-academy (root)", "/", "Scripts, shared lint/format deps, engines"],
            ["@kia-academy/web", "apps/web", "Next.js frontend (port 3000)"],
            ["@kia-academy/api", "apps/api", "NestJS API (port 3001, prefix /api)"],
            ["@kia-academy/shared", "packages/shared", "Shared domain types & algorithms"],
        ],
    )
    add_heading(doc, "3.3 Key directories", 2)
    add_bullets(
        doc,
        [
            "apps/web/src/app — Next.js routes (pages)",
            "apps/web/src/components — shared UI (layout, exam, wizard, admin pieces)",
            "apps/web/src/features/material — Material Studio (data/utils/panels/controller)",
            "apps/web/src/styles — base.css tokens, admin.css, feature CSS",
            "apps/web/src/i18n — fa (default RTL), en, de, es message catalogs",
            "apps/api/src — Nest modules (auth, assessments, readiness, payments, admin, …)",
            "apps/api/prisma — schema.prisma + migrations + seed",
            "packages/shared/src — exam/, constants/, types/, utils/",
            "docs/ — architecture & admin settings catalog",
            "docker/, scripts/, .github/workflows/ — deploy & ops",
        ],
    )
    add_heading(doc, "3.4 Build order", 2)
    add_p(
        doc,
        "Always build @kia-academy/shared before api seed/build or you hit MODULE_NOT_FOUND for "
        "@kia-academy/shared/dist. Root `pnpm build` runs shared → api → web.",
    )

    # ── 4. Languages ───────────────────────────────────────────────────────
    add_heading(doc, "4. Languages & Core Technologies", 1)
    add_table(
        doc,
        ["Technology", "Version / note", "Where used"],
        [
            ["TypeScript", "^5.8.3", "web, api, shared"],
            ["JavaScript", "Config / tooling ESM", "eslint, scripts"],
            ["Node.js", ">=22.13 (.nvmrc 22)", "Runtime for Next & Nest"],
            ["React", "^19.1.0", "apps/web UI"],
            ["Next.js", "^15.5.21", "App Router; /api rewrite; static export for Pages"],
            ["NestJS", "^11.0.12", "apps/api HTTP API"],
            ["Prisma", "^6.19.3", "ORM + migrations"],
            ["PostgreSQL", "16", "Primary datastore"],
            ["SQL", "Via Prisma migrations", "apps/api/prisma/migrations"],
            ["CSS", "Plain CSS + CSS variables (no Tailwind)", "apps/web/src/styles"],
            ["HTML / JSX / TSX", "—", "UI + Material Studio HTML strings"],
            ["JSON", "Settings & blobs", "SiteSetting, assessment answers, adminPanelAccess"],
            ["Shell / PowerShell", "—", "docker entrypoints, portable scripts"],
            ["Markdown", "Lesson content", "Rendered in lesson player"],
        ],
    )
    add_p(
        doc,
        "Not used: Tailwind, Redux, React Query, Turbo/Nx, styled-components, husky, VS Code extensions recommendations file.",
        italic=True,
    )

    # ── 5. Tools & libraries ───────────────────────────────────────────────
    add_heading(doc, "5. Tools, Libraries, Frameworks & Tooling", 1)
    add_heading(doc, "5.1 Package manager & workspace", 2)
    add_bullets(
        doc,
        [
            "pnpm 11.13.0 via Corepack (packageManager field); workspace apps/* + packages/*",
            "concurrently — run web + api together in `pnpm dev`",
            "No Turbo/Nx — root scripts orchestrate builds",
        ],
    )
    add_heading(doc, "5.2 Frontend libraries (@kia-academy/web)", 2)
    add_table(
        doc,
        ["Library", "Purpose"],
        [
            ["next", "App Router, routing, rewrite proxy, static export / standalone"],
            ["react / react-dom", "UI runtime"],
            ["lucide-react", "Icon set across learner + admin UI"],
            ["@kia-academy/shared", "Domain types, exam grading helpers, tracks, prices"],
            ["@playwright/test (dev)", "End-to-end browser tests"],
            ["vitest (dev)", "Unit tests for web"],
            ["typescript (dev)", "Typechecking"],
        ],
    )
    add_heading(doc, "5.3 Backend libraries (@kia-academy/api)", 2)
    add_table(
        doc,
        ["Library", "Purpose"],
        [
            ["@nestjs/common|core|platform-express", "Nest HTTP application"],
            ["@nestjs/config", "Environment configuration"],
            ["@nestjs/jwt + @nestjs/passport + passport-jwt", "Access & refresh JWT auth"],
            ["@nestjs/throttler", "Global and endpoint rate limits"],
            ["@prisma/client + prisma", "Database access & schema"],
            ["bcrypt", "Password hashing (12 rounds)"],
            ["class-validator + class-transformer", "DTO validation / transform"],
            ["cookie-parser", "HttpOnly refresh token cookies"],
            ["helmet", "HTTP security headers"],
            ["joi", "Startup env schema validation"],
            ["multer", "Lesson video uploads (memory + size limits)"],
            ["nodemailer", "Optional SMTP; otherwise console + EmailLog"],
            ["stripe", "Optional Stripe payment provider"],
            ["rxjs / reflect-metadata", "Nest runtime"],
            ["jest + ts-jest (dev)", "API unit tests"],
            ["tsx / ts-node (dev)", "Scripts / seed execution"],
        ],
    )
    add_heading(doc, "5.4 Shared package", 2)
    add_p(
        doc,
        "@kia-academy/shared is pure TypeScript (no runtime deps). It owns exam question bank, grading, "
        "outcome mutation, track catalog, PRODUCT_PRICES / MODULE_PRICES, auth helpers (phone normalize, "
        "unsafe text checks), entitlements key helpers, and site-settings types. Tested with Vitest.",
    )
    add_heading(doc, "5.5 Quality & formatting", 2)
    add_bullets(
        doc,
        [
            "ESLint 9 flat config (eslint.config.mjs) + typescript-eslint + @eslint/js",
            "Prettier 3 (.prettierrc: semi, singleQuote, trailingComma all, printWidth 100, tabWidth 2)",
            "Strict TypeScript across packages; Next plugin in web tsconfig",
        ],
    )
    add_heading(doc, "5.6 Docker & deployment tooling", 2)
    add_bullets(
        doc,
        [
            "docker-compose.yml — Postgres 16-alpine; optional full profile for api+web",
            "docker-compose.ghcr.yml — pull prebuilt GHCR images",
            "Dockerfile — multi-stage Node 22 + pnpm targets for api/web",
            "GitHub Actions: ci.yml, deploy-github-pages.yml, docker-publish.yml",
            "Dependabot: npm weekly, Actions weekly, Docker monthly",
            "PowerShell helpers under scripts/ for Windows portable zip / docker reset",
        ],
    )
    add_heading(doc, "5.7 Extensions / IDE / skills", 2)
    add_p(
        doc,
        "The repo does not ship .vscode/extensions.json, Cursor skills, or husky/pre-commit hooks. "
        "Recommended locally: Node 22+, pnpm via Corepack, PostgreSQL 16 client tools, "
        "Playwright Chromium for e2e (`pnpm --filter @kia-academy/web exec playwright install chromium`).",
    )

    # ── 6. Roles ───────────────────────────────────────────────────────────
    add_heading(doc, "6. Client Roles & Access Control", 1)
    add_heading(doc, "6.1 Prisma UserRole enum", 2)
    add_table(
        doc,
        ["Role", "Who", "Capabilities (summary)"],
        [
            [
                "Guest (no account)",
                "Anonymous visitor",
                "Landing, Material Studio, contact, legal, login. No TopBar/Footer. Soft-gated learner routes → /education.",
            ],
            [
                "LEARNER",
                "Default after OTP/register",
                "Assessment, readiness, roadmap, checkout, courses, learn, dashboard, bootcamp, rewards (JWT + usually profileComplete).",
            ],
            [
                "ADMIN",
                "Moderator",
                "Admin panel sections allowed by User.adminPanelAccess (or site default template). Email/password login.",
            ],
            [
                "SUPER_ADMIN",
                "Full staff",
                "All admin sections; payment gateway settings; default adminAccess template; promote users; assign SUPER_ADMIN.",
            ],
        ],
    )
    add_heading(doc, "6.2 Admin panel access matrix", 2)
    add_p(
        doc,
        "Sections: stats | settings | courses | challenges | users | payments. "
        "Levels per section: view | manage | edit. "
        "Site settings → Admin access is only the DEFAULT template for newly promoted moderators. "
        "Effective access for an ADMIN is resolved from their per-user JSON (Admin → Users matrix). "
        "API enforces via AdminAccessGuard + @AdminAccess(section, level). Sidebar filters in admin/layout.tsx.",
    )
    add_heading(doc, "6.3 Seed accounts (local)", 2)
    add_table(
        doc,
        ["Email", "Role", "Password"],
        [
            ["admin@kia.academy", "SUPER_ADMIN", "KiaAcademy123!"],
            ["moderator@kia.academy", "ADMIN (matrix A)", "KiaAcademy123!"],
            ["moderator2@kia.academy", "ADMIN (matrix B)", "KiaAcademy123!"],
            ["alex@kia.academy", "LEARNER", "KiaAcademy123!"],
        ],
    )

    # ── 7. Sections / routes ───────────────────────────────────────────────
    add_heading(doc, "7. Website Sections & Routes", 1)
    add_p(doc, "General product surfaces a new developer will navigate:")
    add_table(
        doc,
        ["Route", "Section", "Notes"],
        [
            ["/", "Landing", "Minimal Persian hero; no site header for guests"],
            ["/material", "Material Studio", "Public design tools"],
            ["/education", "Education onboarding", "Phone → OTP → profile → start"],
            ["/assessment", "Goal wizard + exam", "Requires complete profile"],
            ["/readiness", "Readiness gate", "Intro / retake entry"],
            ["/readiness/test", "Live exam", "Timed ExamPlayer"],
            ["/readiness/results", "Scorecard", "Then continue to roadmap"],
            ["/roadmap", "Personalized path", "Tree + purchase CTAs"],
            ["/checkout (+ gateway/success/cancel)", "Payments", "Review → PSP → result"],
            ["/courses", "Course catalog", "Enroll / continue"],
            ["/learn/...", "Lesson player", "Video + notes + playground"],
            ["/dashboard", "Learner hub", "Tiles + test history"],
            ["/bootcamp (+ /challenge)", "Arena", "Leaderboard + FizzBuzz"],
            ["/rewards", "Unlocks", "Bonus courses after challenge"],
            ["/contact", "Contact form", "Public"],
            ["/privacy, /terms", "Legal", "Static"],
            ["/login", "Email auth", "Staff / seed learners"],
            ["/register", "Alias", "Redirects to /education"],
            ["/admin/*", "Admin panel", "Stats, users, courses, challenges, finance, payments, analytics, settings, contact inbox"],
        ],
    )
    add_p(
        doc,
        "Chrome rule: SiteChrome renders TopBar + Footer only when profileComplete and not under /admin. "
        "DemoBanner may show when demo mode is on (non-admin).",
    )

    # ── 8. Features (LARGE) ────────────────────────────────────────────────
    add_heading(doc, "8. Complete Feature Catalog", 1)
    add_p(
        doc,
        "This section lists product features down to individual controls where practical. "
        "File paths are under apps/web unless noted.",
    )

    add_heading(doc, "8.1 Global chrome & shared UI", 2)
    add_bullets(
        doc,
        [
            "DemoBanner — warns visitors they are in demo/mock mode",
            "TopBar logo — SUPER_ADMIN → /admin; else dashboard or education",
            "Mobile nav toggle — open/close navigation",
            "Nav links — Material, Education, Contact, Courses, Dashboard (learner/admin); Admin-only for SUPER_ADMIN",
            "LanguageSelector — en / de / es / fa",
            "Theme toggle — light / dark (data-theme)",
            "User chip dropdown — Dashboard, My courses, Rewards, Admin, Sign out",
            "Footer groups — learning, explore, legal deep links",
            "PageBackButton — labeled back navigation",
            "Modal + confetti — challenge/reward completion dialogs",
        ],
    )

    add_heading(doc, "8.2 Landing (/)", 2)
    add_bullets(
        doc,
        [
            "Brand mark کیا آکادمی (hero-level)",
            "Eyebrow + hero headline + short supporting body",
            "CTA → /material (رایگان / free badge)",
            "CTA → /education (start learning)",
            "Footer mini-links: Contact, Privacy, Terms",
        ],
    )

    add_heading(doc, "8.3 Material Studio (/material)", 2)
    add_p(doc, "Shell: brand link home, بازگشت back, tabs پالت / آیکون / انیمیشن / ابزار استایل, toast feedback.")
    add_heading(doc, "Palette tab", 3)
    add_bullets(
        doc,
        [
            "Color search filter; All / None / Light / Core / Dark shade presets",
            "Per-shade chips; matrix swatches (click copy hex); row CSS copy; Copy CSS Variables",
            "Generator: color A/B pickers + hex, HSL/RGB blend, count chips, Random Bases / Reverse / Generate / Copy Palette",
            "Token export format select (CSS/JSON/Tailwind/SCSS); Generate Tokens / Copy Output",
            "Gradient start/end + angle slider; Copy Gradient CSS",
        ],
    )
    add_heading(doc, "Icons tab", 3)
    add_bullets(
        doc,
        [
            "Search; size slider; Copy Favorites Sprite",
            "Per icon: Copy SVG; ★ favorite toggle",
        ],
    )
    add_heading(doc, "Animations tab", 3)
    add_bullets(
        doc,
        [
            "Bezier x1,y1,x2,y2 sliders; Play / Copy cubic-bezier",
            "Animation search; Duration / Delay / Loops; preview text; select card; Copy Full Animation Snippet",
        ],
    )
    add_heading(doc, "Style tools tab", 3)
    add_bullets(
        doc,
        [
            "Font catalog search + size slider; font cards (select/copy)",
            "Contrast text/bg pickers; Suggest Text Color (readable auto-pick)",
            "State persistence via features/material/lib/storage.ts (localStorage)",
        ],
    )

    add_heading(doc, "8.4 Education onboarding (/education)", 2)
    add_heading(doc, "Step phone", 3)
    add_bullets(
        doc,
        [
            "Phone input (normalize to 09xxxxxxxxx)",
            "Submit → POST /api/auth/otp/request → step otp",
            "Back → /",
        ],
    )
    add_heading(doc, "Step otp", 3)
    add_bullets(
        doc,
        [
            "6-digit code field; optional devCode hint in non-prod",
            "Verify → POST /api/auth/otp/verify → session; route to profile or start",
            "Change phone; Resend OTP",
        ],
    )
    add_heading(doc, "Step profile", 3)
    add_bullets(
        doc,
        [
            "Phone disabled read-only; firstName, lastName, city, email",
            "Client + server sanitize unsafe text; Submit → POST /api/auth/profile",
        ],
    )
    add_heading(doc, "Step start", 3)
    add_bullets(
        doc,
        [
            "Continue / Start assessment → ?next= or /assessment",
            "Back → /",
        ],
    )

    add_heading(doc, "8.5 Assessment wizard (/assessment)", 2)
    add_bullets(
        doc,
        [
            "ProgressTrack stage dots",
            "Goal stage — OptCards: job / startup / freelance / fun",
            "Skills — HTML/CSS, JS, Python × Never / Beginner / Comfortable",
            "Personality — sliders solo↔team, structured↔exploratory",
            "Interests — multi-select web, ai, mobile, game, data, backend (≥1)",
            "Learning style — video / reading / building",
            "Hours — slider 3–40 h/week",
            "Back / Continue; final Continue = Start exam → completeWizard() saves assessment+roadmap then ExamPlayer",
        ],
    )

    add_heading(doc, "8.6 Readiness gate, exam & results", 2)
    add_bullets(
        doc,
        [
            "/readiness — domain cover cards (digitalOps, logicalReasoning, techReading, codeSense, problemSolving); Free badge; Start / View results / View roadmap / Retake",
            "/readiness/test — UnifiedTestFlow readinessOnly → api.startExam",
            "ExamPlayer: domain label, Q of N, 30-min countdown (urgent ≤5 min), auto-submit at 0",
            "Dot track jump; answered styling WITHOUT correctness colors",
            "Question types: single_choice radios; multi_choice checkboxes; order ↑↓; fill_blank inputs",
            "Prev / Next / Submit; domain pills answered/total; debounced PATCH answers",
            "/readiness/results: average %, pass chip, RadarChart, score bars, outcome (unlocks/refreshers), verdict, View roadmap, Back dashboard, ?testId= history",
            "Legacy (not live path): FileExplorerTask, CodeFillTask, ReorderTask, FlowchartTask, EnglishReadinessTask",
        ],
    )

    add_heading(doc, "8.7 Roadmap (/roadmap)", 2)
    add_bullets(
        doc,
        [
            "Profile chips: goal, level, style, hours/week",
            "RoadmapTree nodes: completed / up next / locked",
            "Course picker checkboxes + unit prices",
            "Purchase selected courses → checkout COURSE&slugs=",
            "Bundle card: per-module prices, strike/discount, feature list",
            "Enroll bundle → entitlement check or checkout ROADMAP_BUNDLE&roadmapId=",
        ],
    )

    add_heading(doc, "8.8 Checkout & payments UI", 2)
    add_bullets(
        doc,
        [
            "Review: product summary, buyer/email/product/amount/provider read-only",
            "Pay → api.checkout → redirect checkoutUrl or in-app confirmPayment",
            "Stripe return ?session_id= auto-confirm",
            "Gateway page (Zarinpal/IDPay sandbox UX): amount/product/merchant; Complete payment; Cancel/fail",
            "Success / Cancel landing pages",
            "Admin (SUPER_ADMIN): provider select, displayName, merchantId, apiKey, sandbox checkbox",
        ],
    )

    add_heading(doc, "8.9 Courses & Learn player", 2)
    add_bullets(
        doc,
        [
            "/courses cards: icon, title, description, lesson count, progress; Enroll; Continue",
            "Learn sidebar: course title, lesson search, progress bar, lesson links (duration + completed)",
            "Main: back to courses, title/duration/completed chip",
            "LessonVideo: HTML5 controls, fullscreen, placeholder without URL",
            "Markdown content render; Previous / Next / All lessons; Mark complete",
            "Notes textarea (localStorage kia-lesson-note:…); Copy notes",
            "Playground: Auto-run; Run / Reset; preview sizes fluid/desktop/tablet/mobile; HTML/CSS/JS tabs; persisted code; sandboxed iframe; Console + Clear",
        ],
    )

    add_heading(doc, "8.10 Dashboard, Bootcamp, Rewards", 2)
    add_bullets(
        doc,
        [
            "Dashboard empty tiles: Start assessment / Browse courses",
            "Hub tiles: Roadmap, Test results, Bootcamp, Courses, Next lesson; test history list",
            "Bootcamp: rank card, ChallengeCard FizzBuzz (timer → challenge), Palindrome disabled, leaderboard",
            "Challenge page: 15:00 timer, problem panel, code textarea, live scoreFizzBuzz %, Submit → confetti modal → /rewards",
            "Rewards: top-3 unlocked, Interview course unlock/link, Mentor locked, My courses, CTA dashboard",
        ],
    )

    add_heading(doc, "8.11 Auth login, Contact, Legal", 2)
    add_bullets(
        doc,
        [
            "/login: email, password (≥8), submit with ?next=, link toward education",
            "/contact: name, email, subject, message + support aside; POST /api/contact (throttled)",
            "/privacy, /terms: static policy pages",
        ],
    )

    add_heading(doc, "8.12 Admin panel features", 2)
    add_heading(doc, "Shell", 3)
    add_bullets(
        doc,
        [
            "Mobile menu; filtered sidebar by access; Back to site; role badge",
        ],
    )
    add_heading(doc, "Stats (/admin)", 3)
    add_bullets(
        doc,
        [
            "Stat cards: users, revenue, payments, active challenges",
            "7-day revenue chart; snapshot courses/lessons/enrollments/challenges; recent payments; quick links",
        ],
    )
    add_heading(doc, "Users", 3)
    add_bullets(
        doc,
        [
            "List: search, role filter, table, role select + save, AdminAccessMatrix + save for ADMIN, Create user link",
            "Create: name, email, phone, password, role (SUPER_ADMIN only for SUPER_ADMIN role)",
            "Roles page: documentation / matrix overview",
        ],
    )
    add_heading(doc, "Courses admin", 3)
    add_bullets(
        doc,
        [
            "List: New / Edit / Delete (confirm)",
            "New/Edit course fields: title, slug, description, icon, track, sort, published",
            "Lessons: Edit/Delete; form slug/title/content/duration/sort; video upload + preview + remove",
        ],
    )
    add_heading(doc, "Challenges admin", 3)
    add_bullets(
        doc,
        [
            "Form: slug, title, description, points, startsAt, endsAt, active, starterCode; Create/Update; Edit/Delete list",
        ],
    )
    add_heading(doc, "Finance & Payments & Analytics", 3)
    add_bullets(
        doc,
        [
            "Finance KPIs revenue/pending/completed; 14-day settlements",
            "Payments table: search + status filter + ledger",
            "Analytics: enrollment rate, 6-month revenue, status mix, revenue by product type",
        ],
    )
    add_heading(doc, "Settings tabs (/admin/settings)", 3)
    add_bullets(
        doc,
        [
            "General: siteName, tagline, heroMinutes, heroRoadmapsCount, heroMatchPercent, supportEmail",
            "Pricing: course price (toman UI), bundleDiscountPercent, modulePrices Add/Remove",
            "Payment (super): provider, displayName, merchantId, apiKey, sandbox",
            "Tracks: TracksEditor CRUD (key, name, icon, description, modules)",
            "Readiness: passThreshold, pass/fail titles & messages",
            "Bootcamp: unlockScoreThreshold, unlockCourseSlug, defaultRank, defaultPoints",
            "Courses: inline course table",
            "Admin access (super): default matrix template + Save",
            "Backup (super): Export settings JSON",
            "Save section button → PUT /api/admin/settings",
        ],
    )
    add_heading(doc, "Contact inbox", 3)
    add_bullets(
        doc,
        [
            "Table date/from/subject/message/status; Mark read button",
        ],
    )

    add_heading(doc, "8.13 Backend capability features (non-UI)", 2)
    add_bullets(
        doc,
        [
            "JWT access + refresh rotation with DB-backed RefreshToken rows",
            "Phone OTP hashed at rest (SHA-256 of phone:code), TTL, attempt caps",
            "Entitlements grant after successful payment confirmation",
            "Enrollment + LessonProgress tracking",
            "EmailService transactional templates logged to EmailLog",
            "Stripe webhook endpoint for provider=stripe",
            "Authenticated media URLs for lesson videos (not public static dumps)",
            "Health check GET /api/health",
            "Public GET /api/settings for client site configuration",
        ],
    )

    # ── 9. Colors ──────────────────────────────────────────────────────────
    add_heading(doc, "9. Color Palette", 1)
    add_heading(doc, "9.1 Initially requested colors", 2)
    add_p(
        doc,
        "The following hex values were requested for inclusion. A full-repo search found ZERO matches for them "
        "in current source. They are documented here as a historical / reference set that is NOT the active Kia theme:",
    )
    add_table(
        doc,
        ["Hex", "Status in repo", "Notes"],
        [
            ["#6464FF", "Not present", "Legacy indigo-like accent — replaced by Kia teal tokens"],
            ["#C4EED9", "Not present", "Legacy mint — no CSS variable uses this"],
            ["#FFC864", "Not present", "Legacy amber — current gold is #c9a959 / #d4b86a"],
            ["#0E1626", "Not present", "Legacy near-black — current deep is #0a2f44 / dark bg #071821"],
        ],
    )

    add_heading(doc, "9.2 Active Kia brand tokens (light) — styles/base.css", 2)
    add_table(
        doc,
        ["Token", "Hex / value", "Role"],
        [
            ["--kia-deep", "#0a2f44", "Deep brand / fills"],
            ["--kia-accent", "#1f6e8c", "Primary accent"],
            ["--kia-soft", "#2e8a99", "Soft teal"],
            ["--kia-gold", "#c9a959", "Gold accent / amber alias"],
            ["--bg", "#e8f1f6", "Page background"],
            ["--bg-elevated", "#f7fbfd", "Elevated surfaces"],
            ["--text", "#0f1e2c", "Primary text"],
            ["--text-dim", "#385e77", "Secondary text"],
            ["--text-faint", "#5d7f95", "Muted text"],
            ["--emerald / fill", "#1a7a5c", "Success"],
            ["--danger", "#c0392b", "Errors / destructive"],
            ["--on-fill", "#ffffff", "Text on filled buttons"],
            ["--code-bg", "#0a2f44", "Code block background"],
            ["--code-bg-elevated", "#123a52", "Code elevated"],
            ["--code-text", "#e8f1f6", "Code foreground"],
            ["--code-muted", "#9bb8c9", "Code muted"],
            ["--code-keyword", "#7ec8de", "Syntax keyword"],
            ["--code-function", "#e2c77a", "Syntax function"],
            ["--medal-silver", "#6b8496", "Medal UI"],
            ["--medal-bronze", "#9a6b2f", "Medal UI"],
            ["Hero gradient stops", "#e9f0fa, #d6e3f0", "Atmosphere"],
            ["Cards / glass", "rgba(255,255,255,0.88/0.72)", "Surfaces"],
            ["Borders", "rgba(31,110,140,0.22/0.12)", "Dividers"],
        ],
    )

    add_heading(doc, "9.3 Dark theme tokens — [data-theme='dark']", 2)
    add_table(
        doc,
        ["Token", "Hex", "Role"],
        [
            ["--kia-deep", "#7ec8de", "Lightened deep for dark UI"],
            ["--kia-accent", "#3d9bb5", "Accent"],
            ["--kia-soft", "#2e8a99", "Soft teal"],
            ["--kia-gold", "#d4b86a", "Gold"],
            ["--bg", "#071821", "Page background"],
            ["--bg-elevated", "#0d2433", "Elevated"],
            ["--bg-card", "#123044", "Cards"],
            ["--text", "#e8f1f6", "Primary text"],
            ["--text-dim", "#9bb8c9", "Secondary"],
            ["--text-faint", "#6f8fa3", "Muted"],
            ["--indigo / bright", "#7ec8de / #a8dced", "Legacy aliases"],
            ["--teal", "#5ec4c8", "Teal"],
            ["--emerald", "#4ade80", "Success"],
            ["--danger", "#f0a0a0", "Danger"],
            ["--indigo-fill", "#1f6e8c", "Filled controls"],
            ["--teal-fill", "#0a2f44", "Filled controls"],
            ["--emerald-fill", "#047857", "Success fill"],
            ["--medal-bronze", "#c9a959", "Medal"],
        ],
    )

    add_heading(doc, "9.4 Admin panel palette — styles/admin.css", 2)
    add_bullets(
        doc,
        [
            "Backgrounds: #0f0f23, #1a1a2e, #16213e, #12122a, #1a1400",
            "Accents/gold: #ffd700, #ffed4e, #b7791f, #fbbf24",
            "Text/neutrals: #f3f4f6, #9ca3af, #000, #fff",
            "Status: #f87171 (danger), #4ade80 (ok), #60a5fa (info)",
        ],
    )

    add_heading(doc, "9.5 Other hardcoded UI colors", 2)
    add_table(
        doc,
        ["Hex", "Location", "Use"],
        [
            ["#0b0f19, #c9d7e2", "catalog-lesson.css", "Video chrome"],
            ["#6C63FF, #10B981, #F5A524, #EDEFF7", "components/ui/Modal.tsx", "Modal accents"],
            ["#062030, #7b96a8, #0f172a, #e2e8f0, #c7d6e3, #fda4af, #fcd34d", "material-studio.css", "Studio chrome"],
            ["#1f6e8c, #c9a959, #0f172a, #ffffff", "material/lib/storage.ts", "Studio defaults"],
        ],
    )

    add_heading(doc, "9.6 Material Studio full swatch catalog (palette.ts)", 2)
    add_p(
        doc,
        "Used by Material Studio only (not site chrome). Families include Material Design Red→BlueGrey "
        "(shades 50–900) plus custom Kia1 and Kia2 families (~210 hex entries total).",
    )
    add_bullets(
        doc,
        [
            "Families: Red, Pink, Purple, DeepPurple, Indigo, Blue, LightBlue, Cyan, Teal, Green, "
            "LightGreen, Lime, Yellow, Amber, Orange, DeepOrange, Brown, Grey, BlueGrey, Kia1, Kia2",
            "Kia1 samples: #F7F4EF, #F3E4A0, #B8A99A, #C47E5A, #A47764, #5C5E60, #6A3B4C, #2A6B6F, #39615C, #1F4A87",
            "Kia2 samples: #F8F9FA, #C5C6C7, #D4A5A5, #F3A712, #CC5500, #2E8B57, #6A4C93, #2B4C7E, #1F2833, #0B0C10",
            "Standard Material 500 examples: Red #F44336, Pink #E91E63, Purple #9C27B0, Indigo #3F51B5, "
            "Blue #2196F3, Cyan #00BCD4, Teal #009688, Green #4CAF50, Amber #FFC107, Grey #9E9E9E, BlueGrey #607D8B",
        ],
    )
    add_p(
        doc,
        "Viewport theme-color meta uses #e8f1f6 (light) / #071821 (dark) in app/layout.tsx.",
        italic=True,
    )

    # ── 10. Fonts ──────────────────────────────────────────────────────────
    add_heading(doc, "10. Typography & Fonts", 1)
    add_heading(doc, "10.1 Site fonts (next/font/google in app/layout.tsx)", 2)
    add_table(
        doc,
        ["Font", "CSS variable", "Weights", "Role"],
        [
            ["Inter", "--font-inter", "400–800", "LTR body (--font-body)"],
            ["JetBrains Mono", "--font-jetbrains", "400–600", "Code / mono (--font-mono)"],
            ["Vazirmatn", "--font-vazirmatn", "400–800 (arabic subset)", "RTL / Persian UI body"],
        ],
    )
    add_p(
        doc,
        "Fallbacks: Inter → system-ui → sans-serif; JetBrains Mono → monospace; "
        "RTL: Vazirmatn → Tahoma → Arial → sans-serif. Admin inherits (--admin-font: inherit).",
    )
    add_heading(doc, "10.2 System / editor monospace", 2)
    add_bullets(
        doc,
        [
            "Material Studio / controllers may use ui-monospace, SFMono-Regular, Consolas, Liberation Mono, Menlo",
            "Lesson playground sample CSS may reference system-ui",
        ],
    )
    add_heading(doc, "10.3 Material Studio Google Font catalog (on-demand)", 2)
    add_p(
        doc,
        "Loaded via fonts.googleapis.com when selected in Style tools — not part of global site chrome:",
    )
    add_bullets(
        doc,
        [
            "Poppins, Manrope, Montserrat, Lora, Merriweather, Nunito, Cabin, Raleway, Fira Sans, "
            "IBM Plex Sans, Playfair Display, Source Sans 3, Ubuntu, Barlow, Rubik, Work Sans, "
            "DM Sans, Karla, Quicksand, Space Grotesk",
            "Inter may also be imported via CDN inside MaterialController for studio previews",
        ],
    )

    # ── 11. Algorithms ─────────────────────────────────────────────────────
    add_heading(doc, "11. Algorithms & Data Structures", 1)
    add_table(
        doc,
        ["Name", "Behavior", "Primary location"],
        [
            [
                "Skill level scoring",
                "Sum Never=0, Beginner=1, else=2 → bands absoluteBeginner (≤1), confidentBeginner (≤4), earlyIntermediate",
                "packages/shared/.../roadmap.ts",
            ],
            [
                "Roadmap generation",
                "Primary interest → track modules; price = sum(MODULE_PRICES) × (1 − bundleDiscount)",
                "roadmap.ts + constants/tracks.ts",
            ],
            [
                "Exam grading",
                "Per question: single / multi (set equality) / order (sequence) / fill_blank (trim+lower); domain Map averages; pass ≥ 60%",
                "packages/shared/src/exam/grade.ts",
            ],
            [
                "Fisher–Yates shuffle",
                "Shuffle options/order; seeded RNG from attempt id for stable refresh",
                "grade.ts + readiness.service.ts",
            ],
            [
                "Exam outcome mutation",
                "Pass → unlock first module; ≥85% + absoluteBeginner → promote level; fail → prepend refreshers for domains <50%",
                "exam/outcome.ts",
            ],
            [
                "Question bank",
                "Static EXAM_QUESTION_BANK (~76 ids); blueprint kia-readiness-v2; duration 30 min",
                "exam/bank.ts, exam/types.ts",
            ],
            [
                "Legacy readiness %",
                "Per-module correct/total → average of 5 modules; passThreshold default 60",
                "utils/readiness.ts",
            ],
            [
                "FizzBuzz challenge scorer",
                "Heuristic +25 each for %3, %5, Fizz+Buzz, structure/length checks",
                "utils/challenge.ts",
            ],
            [
                "OTP generation",
                "randomInt 100000–999999; SHA-256 hash phone:code; TTL 5m; max 5 attempts",
                "apps/api/.../auth.service.ts",
            ],
            [
                "Iranian phone normalize",
                "Persian/Arabic digits → ASCII; +98/98 → 09xxxxxxxxx",
                "packages/shared/.../auth.ts",
            ],
            [
                "Entitlement keys",
                "resourceType:resourceId string keys",
                "utils/entitlements.ts",
            ],
            [
                "HTML escape / unsafe text",
                "XSS escape; heuristic blocklist without catastrophic backtracking",
                "utils/escape.ts, types/auth.ts",
            ],
        ],
    )
    add_heading(doc, "11.1 Structures commonly used", 2)
    add_bullets(
        doc,
        [
            "Map — domain scores, question lookup",
            "Set — refresher dedupe, multi-choice equality",
            "Sorted arrays — set-equality compares",
            "JSON string / Json columns — assessment answers, roadmap blobs, adminPanelAccess, site settings",
            "Ordered wizard stage keys — WIZARD_STAGES",
            "Track catalog Record — 6 tracks × 5 modules + MODULE_PRICES IRR array",
        ],
    )

    # ── 12. API ────────────────────────────────────────────────────────────
    add_heading(doc, "12. API Surface", 1)
    add_p(
        doc,
        "Global prefix /api. Nest modules registered in apps/api/src/app.module.ts. "
        "Web rewrites /api to Nest in local/dev; production may use NEXT_PUBLIC_API_URL.",
    )
    add_heading(doc, "12.1 Auth flows", 2)
    add_numbered(
        doc,
        [
            "Learner: POST /auth/otp/request → SMS/dev → POST /auth/otp/verify → JWT + refresh cookie → POST /auth/profile",
            "Staff/seed: POST /auth/register or POST /auth/login (email+password)",
            "Session: GET /auth/me, POST /auth/refresh, POST /auth/logout",
        ],
    )
    add_heading(doc, "12.2 Endpoint map", 2)
    add_table(
        doc,
        ["Area", "Base", "Key methods"],
        [
            ["Auth", "/auth", "register, login, otp/request, otp/verify, profile, refresh, logout, me"],
            ["Health", "/health", "GET /"],
            ["Assessments", "/assessments", "POST /, GET :id"],
            ["Roadmaps", "/roadmaps", "POST /, GET :id, POST :id/enroll"],
            ["Readiness", "/readiness", "exam/start, PATCH exam/:id/answers, POST exam/:id/submit, GET exam/:id, CRUD legacy tests"],
            ["Payments", "/payments", "checkout, confirm/:id, my, :id, webhook"],
            ["Courses", "/courses", "list, :slug, lesson, enroll, complete lesson"],
            ["Challenges", "/challenges", "POST submit, GET :id"],
            ["Bootcamp", "/bootcamp", "leaderboard (public), state (JWT)"],
            ["Contact", "/contact", "POST /"],
            ["Settings", "/", "GET settings; GET/PUT admin/settings"],
            ["Media", "/media", "GET lessons/:lessonId/:filename"],
            ["Admin", "/admin", "stats, courses/lessons CRUD + video, challenges, users/role/access, contact, payments"],
        ],
    )
    add_heading(doc, "12.3 Cross-cutting API behavior", 2)
    add_bullets(
        doc,
        [
            "ValidationPipe: whitelist + forbidNonWhitelisted + transform",
            "Guards: JwtAuthGuard, RolesGuard (SUPER_ADMIN ⊇ ADMIN), AdminAccessGuard, ProfileCompleteGuard",
            "Throttler: global 100/min; OTP request 5/min; OTP verify 10/min; contact 5/min",
            "CORS: CORS_ORIGIN with credentials; Helmet with CORP cross-origin",
            "Optional TRUST_PROXY for reverse proxies",
        ],
    )

    # ── 13. Security ───────────────────────────────────────────────────────
    add_heading(doc, "13. Security Measures", 1)
    add_table(
        doc,
        ["Area", "What is done", "Brief why"],
        [
            ["Password hashing", "bcrypt cost 12", "Slow hash resists brute force"],
            ["Access JWT", "Bearer Authorization; JWT_SECRET; ~15m TTL", "Short-lived API credentials"],
            ["Refresh JWT", "HttpOnly cookie (or Bearer); separate secret; ~7d; DB rows", "Revocable long session"],
            ["Cookie flags", "httpOnly; secure in prod; sameSite strict/lax/none as needed", "Reduce XSS/CSRF exposure"],
            ["OTP", "Hashed at rest; TTL; attempt limits; throttle; consume-on-use", "Prevent OTP stuffing"],
            ["Role guards", "RolesGuard + AdminAccess matrix", "Least privilege for moderators"],
            ["Profile gate", "ProfileCompleteGuard on learner flow APIs", "Force registration completeness"],
            ["Rate limiting", "Nest Throttler global + hot endpoints", "Abuse resistance"],
            ["Helmet + Next headers", "X-Frame-Options DENY, nosniff, Referrer-Policy, Permissions-Policy", "Browser hardening"],
            ["CORS", "Single configured origin + credentials", "Block random browser origins"],
            ["DTO validation", "class-validator + Joi env schema", "Reject malformed input / bad config"],
            ["Upload hardening", "Auth media URLs; multer memory + size caps", "Avoid open static video dumps"],
            ["Input sanitization", "sanitizeProfileText / containsUnsafeText / phone+email validators", "XSS & spam heuristics"],
            ["Secrets", "Env-only JWT_*, DATABASE_URL, Stripe/SMTP", "No secrets in git"],
            ["CSRF note", "No dedicated CSRF token middleware; pattern is SameSite cookie + Bearer JWT for mutations", "Know the tradeoff"],
            ["Dev OTP", "devCode / OTP_DEV_EXPOSE only outside production", "Local DX without SMS"],
        ],
    )
    add_p(
        doc,
        "SECURITY.md in the repo is still a generic GitHub template placeholder; treat this guidebook section "
        "and the Nest main.ts / auth modules as the practical security map.",
        italic=True,
    )

    # ── 14. Database ───────────────────────────────────────────────────────
    add_heading(doc, "14. Database Models (Prisma)", 1)
    add_heading(doc, "14.1 Enums", 2)
    add_bullets(
        doc,
        [
            "UserRole: LEARNER | ADMIN | SUPER_ADMIN",
            "PaymentStatus — lifecycle of Payment rows",
            "ProductType — READINESS_TEST | ROADMAP_BUNDLE | COURSE",
            "EntitlementSource — how access was granted",
        ],
    )
    add_heading(doc, "14.2 Models", 2)
    add_bullets(
        doc,
        [
            "User — identity, role, profileComplete, adminPanelAccess Json",
            "PhoneOtp — hashed OTP challenges",
            "RefreshToken — refresh session rows",
            "Assessment — wizard answers blob",
            "Roadmap — personalized path + pricing snapshot",
            "ReadinessTest — legacy/module style results",
            "ExamAttempt — timed exam state/answers/scores",
            "BootcampProfile — arena state",
            "ChallengeSubmission — challenge attempts",
            "Payment — checkout ledger",
            "Entitlement — unlocks",
            "Course / Lesson — catalog",
            "Enrollment / LessonProgress — learner progress",
            "Challenge — bootcamp challenges",
            "EmailLog — outbound email audit",
            "SiteSetting — JSON site configuration",
            "ContactMessage — contact inbox",
        ],
    )
    add_p(
        doc,
        "Prefer `pnpm db:migrate:deploy` in non-interactive CI/agents. `prisma migrate dev` needs CREATEDB "
        "on the kia-academy role for shadow DB (otherwise Prisma P3014).",
    )

    # ── 15. Payments ───────────────────────────────────────────────────────
    add_heading(doc, "15. Payments, Currency & Email", 1)
    add_bullets(
        doc,
        [
            "Currency: amounts stored as IRR integers; Persian UI shows تومان (typically ÷10)",
            "Catalog: PRODUCT_PRICES and MODULE_PRICES in @kia-academy/shared",
            "Providers: dev | zarinpal | idpay | stripe (SitePaymentSettings)",
            "Checkout is a review step; confirm redirects to configured provider UX",
            "Stripe webhook supported when provider=stripe",
            "Email: Nodemailer SMTP optional; otherwise console logging + EmailLog persistence "
            "(welcome, purchase, readiness results, etc.)",
        ],
    )

    # ── 16. i18n & theme ───────────────────────────────────────────────────
    add_heading(doc, "16. Internationalization & Theming", 1)
    add_bullets(
        doc,
        [
            "Default locale fa with dir=rtl; Vazirmatn for Persian body text",
            "Additional message catalogs: en, de, es (LanguageSelector)",
            "LTR isolate helpers for code, emails, timers, prices (unicode-bidi isolate)",
            "Theme: light default teal/gold; dark via data-theme='dark' CSS variable swap",
            "Motion/transitions on body background/color; product uses intentional CSS motion in hero/studio",
        ],
    )

    # ── 17. Dev setup ──────────────────────────────────────────────────────
    add_heading(doc, "17. Development Setup & Commands", 1)
    add_heading(doc, "17.1 Prerequisites", 2)
    add_bullets(
        doc,
        [
            "Node >= 22.13, Corepack/pnpm 11.13.0",
            "PostgreSQL 16 (Cloud VMs: sudo pg_ctlcluster 16 main start — Docker often unavailable)",
            "Copy env: apps/api/.env.example → .env; apps/web/.env.example → .env.local",
            "Role/DB kia-academy/kia-academy as in .env.example; grant CREATEDB if using migrate dev",
        ],
    )
    add_heading(doc, "17.2 Typical commands", 2)
    add_bullets(
        doc,
        [
            "pnpm install",
            "pnpm --filter @kia-academy/shared build",
            "pnpm db:migrate:deploy",
            "pnpm db:seed",
            "pnpm dev  — web :3000 + api :3001",
            "pnpm lint | typecheck | test | build",
            "pnpm test:e2e — after Playwright Chromium install",
            "pnpm build:pages — static GitHub Pages export with demo mode",
            "pnpm docker:db — only when Docker is available",
        ],
    )

    # ── 18. Testing & CI ───────────────────────────────────────────────────
    add_heading(doc, "18. Testing & CI/CD", 1)
    add_table(
        doc,
        ["Layer", "Tool", "Notes"],
        [
            ["Web unit", "Vitest", "apps/web/vitest.config.ts"],
            ["Shared unit", "Vitest", "packages/shared"],
            ["API unit", "Jest + ts-jest", "apps/api/jest.config.js"],
            ["E2E", "Playwright Chromium", "apps/web/e2e/"],
            ["CI", "GitHub Actions ci.yml", "Lint/test/build + Postgres 16 service"],
            ["Pages deploy", "deploy-github-pages.yml", "Static export"],
            ["Images", "docker-publish.yml", "GHCR"],
        ],
    )

    # ── 19. Seed / demo / ops ──────────────────────────────────────────────
    add_heading(doc, "19. Seed Data, Demo Mode & Ops Notes", 1)
    add_bullets(
        doc,
        [
            "Seed creates SUPER_ADMIN, two ADMIN moderators with different matrices, LEARNER alex, sample courses/lessons (some with videoUrl demos)",
            "NEXT_PUBLIC_DEMO_MODE=true (Pages): mock data in browser; progress local only",
            "OTP in non-production returns/logs devCode",
            "Env files are git-ignored — always copy from examples",
            "postinstall runs prisma generate",
            "pnpm overrides pin fast-uri, postcss, sharp; allowBuilds for prisma/bcrypt/esbuild/sharp",
        ],
    )

    # ── 20. Glossary ───────────────────────────────────────────────────────
    add_heading(doc, "20. Glossary & Quick Reference", 1)
    add_heading(doc, "20.1 Terms", 2)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Kia Academy / کیا آکادمی", "Product name for the Persian-first rebuild"],
            ["@kia-academy/*", "Monorepo package scope"],
            ["Material Studio", "Public design toolkit at /material"],
            ["Education path", "OTP → profile → assessment → exam → roadmap → pay → learn"],
            ["Readiness / preparations test", "Free timed exam after assessment"],
            ["Roadmap", "Personalized module path + bundle pricing"],
            ["profileComplete", "User finished required profile; unlocks chrome + learner APIs"],
            ["adminPanelAccess", "Per-moderator JSON matrix"],
            ["adminAccess", "Site default template for new moderators"],
            ["IRR / تومان", "Stored rials vs displayed tomans"],
            ["ExamPlayer", "Live MCQ/order/fill UI without live correctness colors"],
        ],
    )
    add_heading(doc, "20.2 File-path index", 2)
    add_bullets(
        doc,
        [
            "Architecture: docs/REBUILD_ARCHITECTURE.md",
            "Admin settings catalog: docs/ADMIN_SETTINGS_CATALOG.md",
            "Prisma schema: apps/api/prisma/schema.prisma",
            "Shared domain: packages/shared/src/",
            "Web routes: apps/web/src/app/",
            "Material feature: apps/web/src/features/material/",
            "Design tokens: apps/web/src/styles/base.css",
            "Cloud agent notes: AGENTS.md",
            "README: setup, Pages, business model overview",
        ],
    )
    add_heading(doc, "20.3 How to use this guidebook", 2)
    add_p(
        doc,
        "Start with sections 1–3 for product mental model, section 6–8 when implementing UI, "
        "11–14 when changing scoring or API contracts, 9–10 for design consistency, and 13+17 when "
        "hardening or bootstrapping an environment. Re-run inventory against the repo if features drift.",
    )

    # Footer note
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run(
        "— End of Kia Academy Project Guidebook —\n"
        "Generated from repository analysis for new developer onboarding."
    )
    set_run_font(r, size=9, italic=True, color=RGBColor(0x5D, 0x7F, 0x95))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
